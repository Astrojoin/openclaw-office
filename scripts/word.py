#!/usr/bin/env python3
"""CRUD operations for .docx files — purely offline, no network calls.

Works with bytes in memory (primary) or local file paths (optional).
Uses python-docx for all document manipulation.

Usage:
  from word import WordDoc

  # From bytes (downloaded from OneDrive, etc.)
  doc = WordDoc(data=bytes_data)
  text = doc.read()
  doc.update({"operations": [{"method": "add_paragraph", "args": ["Hello"]}]})
  new_bytes = doc.to_bytes()

  # From local file
  doc = WordDoc(path="/tmp/report.docx")
  doc.update({"operations": [{"method": "add_heading", "args": ["Title", 0]}]})
  doc.save()  # saves back to same path

  # Create blank document
  doc = WordDoc.create()
  doc.update({"operations": [{"method": "add_paragraph", "args": ["New document"]}]})
  new_bytes = doc.to_bytes()
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordDoc:
    """In-memory .docx editor using python-docx."""

    def __init__(self, data: bytes | None = None, path: str | None = None):
        """
        Initialize from bytes (in-memory) or a local file path.
        If both provided, bytes take precedence.
        If neither, raises ValueError (use WordDoc.create() for blank docs).
        """
        self.path = Path(path) if path else None
        self._doc: Document | None = None

        if data is not None:
            self._doc = Document(io.BytesIO(data))
        elif path is not None:
            self._doc = Document(path)
        else:
            raise ValueError("Provide either 'data' (bytes) or 'path' (str), or use WordDoc.create()")

    @classmethod
    def create(cls) -> "WordDoc":
        """Create a blank Word document."""
        instance = cls.__new__(cls)
        instance.path = None
        instance._doc = Document()
        return instance

    @property
    def doc(self) -> Document:
        """Access the underlying python-docx Document object."""
        return self._doc

    # ── READ ────────────────────────────────────────────────────────────

    def read(self, mode: str = "plain") -> str:
        """Read document content.

        Args:
            mode: "plain" for plain text, "structured" for JSON with paragraphs/styles

        Returns:
            Document content as string (plain text or JSON).
        """
        if mode == "structured":
            return self._read_structured()
        return self._read_plain()

    def _read_plain(self) -> str:
        """Extract all text as plain string."""
        lines = []
        for para in self._doc.paragraphs:
            lines.append(para.text)
        # Include tables
        for table in self._doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines)

    def _read_structured(self) -> str:
        """Extract document content as structured JSON string."""
        import json
        paragraphs = []
        for para in self._doc.paragraphs:
            paragraphs.append({
                "index": para._element.getparent().index(para._element) if para._element.getparent() is not None else None,
                "style": para.style.name if para.style else None,
                "text": para.text,
            })
        tables = []
        for table in self._doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        return json.dumps({"paragraphs": paragraphs, "tables": tables}, indent=2, ensure_ascii=False)

    # ── UPDATE ──────────────────────────────────────────────────────────

    def update(self, spec: dict) -> str:
        """Apply operations to the document.

        Args:
            spec: {"operations": [{"method": "...", "args": [...], "kwargs": {...}}]}

        Supported methods map to python-docx Document and paragraph methods:
          Document-level: add_paragraph, add_heading, add_table, add_page_break,
                          add_picture, add_section
          Paragraph-level: (access via "paragraph_index" in kwargs)
            add_run, clear, insert_paragraph_before

        Returns:
            Summary of operations applied.
        """
        operations = spec.get("operations", [])
        results = []

        for op in operations:
            method_name = op.get("method", "")
            args = op.get("args", [])
            kwargs = op.get("kwargs", {})
            result = self._apply_operation(method_name, args, kwargs)
            results.append(result)

        return f"Applied {len(results)} operations: " + "; ".join(results)

    def _apply_operation(self, method: str, args: list, kwargs: dict) -> str:
        """Apply a single operation to the document."""
        # Paragraph-level operations (need paragraph_index)
        if "paragraph_index" in kwargs:
            idx = kwargs.pop("paragraph_index")
            para = self._doc.paragraphs[idx]
            if hasattr(para, method):
                fn = getattr(para, method)
                # Special: add_run returns a Run, we convert to text
                result = fn(*args, **kwargs)
                return f"paragraph[{idx}].{method}"
            return f"ERROR: paragraph has no method '{method}'"

        # Table-level operations
        if "table_index" in kwargs:
            tidx = kwargs.pop("table_index")
            table = self._doc.tables[tidx]

            if method == "add_row":
                row = table.add_row()
                if args and isinstance(args[0], list):
                    for i, val in enumerate(args[0]):
                        if i < len(row.cells):
                            row.cells[i].text = str(val)
                return f"table[{tidx}].add_row"

            if method == "set_cell":
                ridx = kwargs.pop("row_index", 0)
                cidx = kwargs.pop("col_index", 0)
                table.rows[ridx].cells[cidx].text = str(args[0]) if args else ""
                return f"table[{tidx}].cell[{ridx},{cidx}] = '{args[0] if args else ''}'"

            return f"table[{tidx}].{method}"

        # Document-level operations
        doc = self._doc
        if method == "add_paragraph":
            text = args[0] if args else ""
            style = kwargs.get("style", None)
            para = doc.add_paragraph(text, style=style)
            # Handle run formatting via kwargs
            if "bold" in kwargs or "italic" in kwargs or "font_size" in kwargs:
                run = para.runs[0] if para.runs else para.add_run(text)
                if kwargs.get("bold"):
                    run.bold = True
                if kwargs.get("italic"):
                    run.italic = True
                if "font_size" in kwargs:
                    run.font.size = Pt(kwargs["font_size"])
                if "font_color" in kwargs:
                    r, g, b = kwargs["font_color"]
                    run.font.color.rgb = RGBColor(r, g, b)
            return f"add_paragraph('{text[:50]}...')" if len(text) > 50 else f"add_paragraph('{text}')"

        if method == "add_heading":
            text = args[0] if args else ""
            level = args[1] if len(args) > 1 else kwargs.get("level", 1)
            doc.add_heading(text, level=level)
            return f"add_heading('{text}', level={level})"

        if method == "add_table":
            rows = args[0] if len(args) > 0 else kwargs.get("rows", 2)
            cols = args[1] if len(args) > 1 else kwargs.get("cols", 2)
            table = doc.add_table(rows=rows, cols=cols)
            # Optional: populate from data
            if "data" in kwargs:
                for i, row_data in enumerate(kwargs["data"]):
                    if i < len(table.rows):
                        for j, val in enumerate(row_data):
                            if j < len(table.rows[i].cells):
                                table.rows[i].cells[j].text = str(val)
            style = kwargs.get("style", "Table Grid")
            table.style = style
            return f"add_table({rows}x{cols})"

        if method == "add_page_break":
            doc.add_page_break()
            return "add_page_break"

        if method == "add_picture":
            # For in-memory, picture path must be accessible on disk
            path = args[0] if args else kwargs.get("path", "")
            width = kwargs.get("width", None)
            w = Inches(width) if width else None
            doc.add_picture(path, width=w)
            return f"add_picture('{path}')"

        if method == "remove_paragraph":
            idx = args[0] if args else kwargs.get("index", 0)
            if 0 <= idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]
                p._element.getparent().remove(p._element)
                return f"remove_paragraph({idx})"
            return f"ERROR: paragraph index {idx} out of range"

        if method == "replace_text":
            old = args[0] if len(args) > 0 else kwargs.get("old", "")
            new = args[1] if len(args) > 1 else kwargs.get("new", "")
            count = 0
            for para in doc.paragraphs:
                if old in para.text:
                    for run in para.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            count += 1
            return f"replace_text: {count} replacements"

        # Fallback: try to call method on Document directly
        if hasattr(doc, method):
            fn = getattr(doc, method)
            fn(*args, **kwargs)
            return f"{method}()"

        return f"ERROR: unknown method '{method}'"

    # ── DELETE (local only) ─────────────────────────────────────────────

    def delete(self) -> str:
        """Delete local file if path was provided. Cannot delete cloud files."""
        if self.path and self.path.exists():
            self.path.unlink()
            return f"Deleted: {self.path}"
        return "No local file to delete (in-memory document). Use onedrive.py for cloud deletion."

    # ── OUTPUT ──────────────────────────────────────────────────────────

    def to_bytes(self) -> bytes:
        """Serialize document to bytes (for upload to OneDrive, etc.)."""
        buf = io.BytesIO()
        self._doc.save(buf)
        buf.seek(0)
        return buf.read()

    def save(self, path: str | None = None) -> str:
        """Save to local file. If path omitted, saves to original path."""
        target = Path(path) if path else self.path
        if target is None:
            raise ValueError("No path specified. Use to_bytes() for in-memory output.")
        target.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(target))
        return f"Saved: {target}"


# ── CLI ─────────────────────────────────────────────────────────────────────

def main():
    """Quick demo: create a blank doc, add content, show text."""
    doc = WordDoc.create()
    doc.update({"operations": [
        {"method": "add_heading", "args": ["OneDrive Skill Test", 1]},
        {"method": "add_paragraph", "args": ["This document was created by the openclaw-office skill."]},
    ]})
    print("=== Document content ===")
    print(doc.read())
    print(f"\n=== Bytes size: {len(doc.to_bytes())} bytes ===")


if __name__ == "__main__":
    main()
